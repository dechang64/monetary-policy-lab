#!/usr/bin/env python3
"""
Fix all JMP docx files - v2 with better pattern matching
"""

from docx import Document
from lxml import etree
import re
import os

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def m(tag): return f'{{{M_NS}}}{tag}'
def w(tag): return f'{{{W_NS}}}{tag}'

def make_mr(text, italic=True):
    r = etree.Element(m('r'))
    rPr = etree.SubElement(r, m('rPr'))
    if italic:
        sty = etree.SubElement(rPr, m('sty'))
        sty.set(m('val'), 'i')
    t = etree.SubElement(r, m('t'))
    t.text = text
    return r

def make_sub(base_text, sub_text, base_italic=True):
    sSub = etree.Element(m('sSub'))
    e = etree.SubElement(sSub, m('e'))
    e.append(make_mr(base_text, italic=base_italic))
    sub = etree.SubElement(sSub, m('sub'))
    sub.append(make_mr(sub_text, italic=False))
    return sSub

def make_sup_on_element(base_element, sup_text):
    sSup = etree.Element(m('sSup'))
    e = etree.SubElement(sSup, m('e'))
    e.append(base_element)
    sup = etree.SubElement(sSup, m('sup'))
    sup.append(make_mr(sup_text, italic=False))
    return sSup

def make_frac(num_children, den_children):
    f = etree.Element(m('f'))
    fPr = etree.SubElement(f, m('fPr'))
    ctrlPr = etree.SubElement(fPr, m('ctrlPr'))
    num = etree.SubElement(f, m('num'))
    for child in num_children: num.append(child)
    den = etree.SubElement(f, m('den'))
    for child in den_children: den.append(child)
    return f

def make_hat(base_text, sub_text=None):
    acc = etree.Element(m('acc'))
    accPr = etree.SubElement(acc, m('accPr'))
    chr_elem = etree.SubElement(accPr, m('chr'))
    chr_elem.set(m('val'), '̂')
    ctrlPr = etree.SubElement(accPr, m('ctrlPr'))
    e = etree.SubElement(acc, m('e'))
    if sub_text:
        e.append(make_sub(base_text, sub_text))
    else:
        e.append(make_mr(base_text))
    return acc

def make_d(children):
    d = etree.Element(m('d'))
    dPr = etree.SubElement(d, m('dPr'))
    ctrlPr = etree.SubElement(dPr, m('ctrlPr'))
    e = etree.SubElement(d, m('e'))
    for child in children: e.append(child)
    return d

def make_func(name_children, arg_children):
    func = etree.Element(m('func'))
    funcPr = etree.SubElement(func, m('funcPr'))
    ctrlPr = etree.SubElement(funcPr, m('ctrlPr'))
    fName = etree.SubElement(func, m('fName'))
    for child in name_children: fName.append(child)
    e = etree.SubElement(func, m('e'))
    for child in arg_children: e.append(child)
    return func

def build_formula_1():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('LM', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_frac([make_mr('Positive wordsₜ − Negative wordsₜ', italic=False)],
                            [make_mr('Total wordsₜ', italic=False)]))
    return omath

def build_formula_2():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('CB', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_frac([make_mr('Hawkish wordsₜ − Dovish wordsₜ', italic=False)],
                            [make_mr('Total wordsₜ', italic=False)]))
    return omath

def build_formula_3():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = 0.5 × ', italic=False))
    omath.append(make_sub('LM', 't'))
    omath.append(make_mr(' + 0.5 × ', italic=False))
    omath.append(make_sub('CB', 't'))
    return omath

def build_formula_4():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_5():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('R', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_6():
    omath = etree.Element(m('oMath'))
    omath.append(make_mr('W'))
    omath.append(make_mr(' = ', italic=False))
    paren = make_d([make_hat('β', '1'), make_mr(' − ', italic=False), make_hat('β', '2')])
    omath.append(make_frac([make_sup_on_element(paren, '2')],
                            [make_func([make_mr('Var', italic=False)],
                                       [make_hat('β', '1'), make_mr(' − ', italic=False), make_hat('β', '2')])]))
    return omath

def build_formula_7():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('R', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '3'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '4'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_d([make_sub('S', 't'), make_mr(' × ', italic=False), make_sub('FG', 't')]))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_8():
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_mr('ρ'))
    omath.append(make_sub('S', 't−1'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

BUILDERS = [
    (build_formula_1, '1'), (build_formula_2, '2'), (build_formula_3, '3'),
    (build_formula_4, '4'), (build_formula_5, '5'), (build_formula_6, '6'),
    (build_formula_7, '7'), (build_formula_8, '8'),
]

def make_omath_para(omath):
    p = etree.Element(m('oMathPara'))
    pPr = etree.SubElement(p, m('oMathParaPr'))
    jc = etree.SubElement(pPr, m('jc'))
    jc.set(m('val'), 'center')
    p.append(omath)
    return p

def replace_formula(para, omath_para, eq_num):
    elem = para._element
    pPr = elem.find(w('pPr'))
    for child in list(elem):
        if child.tag != w('pPr'):
            elem.remove(child)
    elem.append(omath_para)
    r = etree.SubElement(elem, w('r'))
    rPr = etree.SubElement(r, w('rPr'))
    rFonts = etree.SubElement(rPr, w('rFonts'))
    rFonts.set(w('ascii'), 'Times New Roman')
    rFonts.set(w('hAnsi'), 'Times New Roman')
    t = etree.SubElement(r, w('t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = f'  ({eq_num})'

def detect_formula(text):
    """Detect which formula this is based on text content. Returns builder index or None."""
    # Must be centered and contain =
    if '=' not in text:
        return None
    
    # Formula 1: LMₜ = ... (starts with LM)
    if re.match(r'(\\text)?LM[ₜ_]', text):
        return 0
    # Formula 2: CBₜ = ... (starts with CB)
    if re.match(r'(\\text)?CB[ₜ_]', text):
        return 1
    # Formula 3: Sₜ = 0.5 × ... (starts with S, has 0.5)
    if re.match(r'S[ₜ_]\s*=\s*0\.5', text):
        return 2
    # Formula 6: W = ... (starts with W)
    if re.match(r'W\s*=', text):
        return 5
    # Formula 4: Sₜ = α + β₁ · Target (starts with S, has Target)
    if re.match(r'S[ₜ_]\s*=\s*(\\alpha|α)', text) and 'Target' in text and 'β₃' not in text and 'ρ' not in text and '\\rho' not in text:
        return 3
    # Formula 5: Rₜ = α + β₁ · Target (starts with R, has Target but NOT β₃)
    if re.match(r'R[ₜ_]\s*=\s*(\\alpha|α|\s*\+)', text) and 'Target' in text and 'β₃' not in text and '\\beta_3' not in text:
        return 4
    # Formula 7: Rₜ = α + β₁ · Target + ... + β₃ · S + β₄ (has β₃/β₄)
    if re.match(r'R[ₜ_]\s*=', text) and ('β₃' in text or 'β₄' in text or '\\beta_3' in text or '\\beta_4' in text):
        return 6
    # Formula 8: Sₜ = α + ρS (has ρ)
    if re.match(r'S[ₜ_]\s*=', text) and ('ρ' in text or '\\rho' in text or 'Sₜ₋₁' in text or 'Sₜ-1' in text or 'S_{t-1}' in text):
        return 7
    
    return None

def fix_file(input_path, output_path=None):
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f'{base}_fixed{ext}'
    
    doc = Document(input_path)
    name = os.path.basename(input_path)
    
    # Step 1: Fix formula paragraphs
    formula_fixes = 0
    for i, para in enumerate(doc.paragraphs):
        if para.alignment != 1:
            continue
        text = para.text.strip()
        if not text:
            continue
        
        idx = detect_formula(text)
        if idx is not None:
            builder, eq_num = BUILDERS[idx]
            omath = builder()
            omath_para = make_omath_para(omath)
            replace_formula(para, omath_para, eq_num)
            formula_fixes += 1
            print(f'  ✓ Formula ({eq_num}) at para {i}')
    
    # Step 2: Remove [Table/Figure about here] placeholders
    placeholder_count = 0
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if ('[Table' in text or '[Figure' in text) and 'about here' in text:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)
                placeholder_count += 1
    if placeholder_count:
        print(f'  ✓ Removed {placeholder_count} placeholders')
    
    # Step 3: Fix inline text
    inline_fixes = 0
    replacements = [
        (r'\hatβ₁', 'β̂₁'), (r'\hatβ₂', 'β̂₂'),
        (r'\hat{\beta_1', 'β̂₁'), (r'\hat{\beta_2', 'β̂₂'),
        (r'\hat\beta_1', 'β̂₁'), (r'\hat\beta_2', 'β̂₂'),
        (r'\beta_1', 'β₁'), (r'\beta_2', 'β₂'),
        (r'\beta_3', 'β₃'), (r'\beta_4', 'β₄'),
        (r'\alpha', 'α'), (r'\varepsilon', 'ε'), (r'\rho', 'ρ'),
        (r'\cdot', '·'), (r'\times', '×'),
        (r'\geq', '≥'), (r'\leq', '≤'),
        (r'\textVar', 'Var'), (r'\textCov', 'Cov'),
        (r'\textTarget', 'Target'), (r'\textPath', 'Path'),
        (r'\frac', ''),
        ('̂_1', 'β̂₁'), ('̂_2', 'β̂₂'),
    ]
    
    for para in doc.paragraphs:
        if para.alignment == 1:
            continue
        for run in para.runs:
            if not run.text:
                continue
            original = run.text
            modified = original
            for old, new in replacements:
                modified = modified.replace(old, new)
            if modified != original:
                run.text = modified
                inline_fixes += 1
    if inline_fixes:
        print(f'  ✓ Fixed {inline_fixes} inline text runs')
    
    doc.save(output_path)
    print(f'  → Saved: {output_path}')
    return formula_fixes, placeholder_count, inline_fixes

def main():
    files = [
        '/home/z/my-project/monetary-policy-lab/delivery_v9.9/Beyond_the_Rate_JMP_v9.docx',
        '/home/z/my-project/monetary-policy-lab/docs/Beyond_the_Rate_JMP_v9.docx',
        '/home/z/my-project/monetary-policy-lab/docs/paper_v9_35page.docx',
    ]
    
    total = [0, 0, 0]
    for fpath in files:
        name = os.path.basename(fpath)
        print(f'\n{"="*60}')
        print(f'Fixing: {name}')
        print(f'{"="*60}')
        try:
            f, p, i = fix_file(fpath)
            total[0] += f; total[1] += p; total[2] += i
        except Exception as e:
            print(f'  ERROR: {e}')
            import traceback; traceback.print_exc()
    
    print(f'\n{"="*60}')
    print(f'Done. Total: {total[0]} formulas, {total[1]} placeholders, {total[2]} inline fixes')

if __name__ == '__main__':
    main()
