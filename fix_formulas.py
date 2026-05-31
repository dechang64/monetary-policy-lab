#!/usr/bin/env python3
"""
Fix all formulas in Beyond_the_Rate_JMP_v9.docx:
1. Convert LaTeX text runs to proper Word oMath objects
2. Fix inline text with missing β symbols
3. Add equation numbers
"""

from docx import Document
from docx.shared import Pt
from lxml import etree
import copy
import re
import sys

# Namespaces
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
MC_NS = 'http://schemas.openxmlformats.org/markup-compatibility/2006'

NSMAP = {
    'm': M_NS,
    'w': W_NS,
    'mc': MC_NS,
}

def m(tag): return f'{{{M_NS}}}{tag}'
def w(tag): return f'{{{W_NS}}}{tag}'

# ─── oMath Builder Helpers ───

def make_mr(text, italic=True):
    """Create an m:r (math run) element"""
    r = etree.Element(m('r'))
    rPr = etree.SubElement(r, m('rPr'))
    if italic:
        sty = etree.SubElement(rPr, m('sty'))
        sty.set(m('val'), 'i')
    t = etree.SubElement(r, m('t'))
    t.text = text
    return r

def make_mr_normal(text):
    """Create an m:r with normal (non-italic) style"""
    r = etree.Element(m('r'))
    rPr = etree.SubElement(r, m('rPr'))
    sty = etree.SubElement(rPr, m('sty'))
    sty.set(m('val'), 'p')
    t = etree.SubElement(r, m('t'))
    t.text = text
    return r

def make_sub(base_text, sub_text, base_italic=True):
    """Create m:sSub (subscript) element"""
    sSub = etree.Element(m('sSub'))
    e = etree.SubElement(sSub, m('e'))
    e.append(make_mr(base_text, italic=base_italic))
    sub = etree.SubElement(sSub, m('sub'))
    sub.append(make_mr(sub_text, italic=False))
    return sSub

def make_sup(base_text, sup_text, base_italic=True):
    """Create m:sSup (superscript) element"""
    sSup = etree.Element(m('sSup'))
    e = etree.SubElement(sSup, m('e'))
    e.append(make_mr(base_text, italic=base_italic))
    sup = etree.SubElement(sSup, m('sup'))
    sup.append(make_mr(sup_text, italic=False))
    return sSup

def make_frac(num_children, den_children):
    """Create m:f (fraction) element"""
    f = etree.Element(m('f'))
    fPr = etree.SubElement(f, m('fPr'))
    ctrlPr = etree.SubElement(fPr, m('ctrlPr'))
    num = etree.SubElement(f, m('num'))
    for child in num_children:
        num.append(child)
    den = etree.SubElement(f, m('den'))
    for child in den_children:
        den.append(child)
    return f

def make_hat(base_text, sub_text=None):
    """Create m:acc (accent/hat) element"""
    acc = etree.Element(m('acc'))
    accPr = etree.SubElement(acc, m('accPr'))
    chr_elem = etree.SubElement(accPr, m('chr'))
    chr_elem.set(m('val'), '̂')  # combining circumflex
    ctrlPr = etree.SubElement(accPr, m('ctrlPr'))
    e = etree.SubElement(acc, m('e'))
    if sub_text:
        e.append(make_sub(base_text, sub_text))
    else:
        e.append(make_mr(base_text))
    return acc

def make_nary(lower, upper, child):
    """Create m:nary (summation etc) - not needed for now"""
    nary = etree.Element(m('nary'))
    naryPr = etree.SubElement(nary, m('naryPr'))
    sub = etree.SubElement(nary, m('sub'))
    sub.append(make_mr(lower))
    sup = etree.SubElement(nary, m('sup'))
    sup.append(make_mr(upper))
    e = etree.SubElement(nary, m('e'))
    e.append(child)
    return nary

def make_d(children):
    """Create m:d (delimiter/parentheses) element"""
    d = etree.Element(m('d'))
    dPr = etree.SubElement(d, m('dPr'))
    ctrlPr = etree.SubElement(dPr, m('ctrlPr'))
    e = etree.SubElement(d, m('e'))
    for child in children:
        e.append(child)
    return d

def make_func(name_children, arg_children):
    """Create m:func (function like Var, Cov) element"""
    func = etree.Element(m('func'))
    funcPr = etree.SubElement(func, m('funcPr'))
    ctrlPr = etree.SubElement(funcPr, m('ctrlPr'))
    fName = etree.SubElement(func, m('fName'))
    for child in name_children:
        fName.append(child)
    e = etree.SubElement(func, m('e'))
    for child in arg_children:
        e.append(child)
    return func

# ─── Build each formula ───

def build_formula_1():
    """LM_t = (Positive words_t - Negative words_t) / Total words_t"""
    omath = etree.Element(m('oMath'))
    # LM_t
    omath.append(make_sub('LM', 't'))
    # =
    omath.append(make_mr(' = ', italic=False))
    # fraction
    num = [
        make_mr('Positive words', italic=False),
        make_sub('', 't', base_italic=False),  # This won't work well, use different approach
    ]
    # Actually, let me simplify: use text for the fraction content
    num = [make_mr('Positive wordsₜ − Negative wordsₜ', italic=False)]
    den = [make_mr('Total wordsₜ', italic=False)]
    omath.append(make_frac(num, den))
    return omath

def build_formula_2():
    """CB_t = (Hawkish words_t - Dovish words_t) / Total words_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('CB', 't'))
    omath.append(make_mr(' = ', italic=False))
    num = [make_mr('Hawkish wordsₜ − Dovish wordsₜ', italic=False)]
    den = [make_mr('Total wordsₜ', italic=False)]
    omath.append(make_frac(num, den))
    return omath

def build_formula_3():
    """S_t = 0.5 × LM_t + 0.5 × CB_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = 0.5 × ', italic=False))
    omath.append(make_sub('LM', 't'))
    omath.append(make_mr(' + 0.5 × ', italic=False))
    omath.append(make_sub('CB', 't'))
    return omath

def build_formula_4():
    """S_t = α + β₁ · Target_t + β₂ · Path_t + ε_t  (H1)"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_5():
    """R_t = α + β₁ · Target_t + β₂ · Path_t + ε_t  (H2)"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('R', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_6():
    """W = (β̂₁ − β̂₂)² / Var(β̂₁ − β̂₂)  (Wald)"""
    omath = etree.Element(m('oMath'))
    omath.append(make_mr('W'))
    omath.append(make_mr(' = ', italic=False))
    
    # Fraction: numerator = (β̂₁ − β̂₂)², denominator = Var(β̂₁ − β̂₂)
    # Numerator: (β̂₁ − β̂₂)²
    num = [
        make_d([
            make_hat('β', '1'),
            make_mr(' − ', italic=False),
            make_hat('β', '2'),
        ]),
        make_sup('', '2'),  # This won't work well
    ]
    # Actually, let me build the superscript on the delimiter
    # (β̂₁ − β̂₂)² 
    paren_content = make_d([
        make_hat('β', '1'),
        make_mr(' − ', italic=False),
        make_hat('β', '2'),
    ])
    num = [make_sup_on_element(paren_content, '2')]
    
    # Denominator: Var(β̂₁ − β̂₂)
    den = [
        make_func(
            [make_mr('Var', italic=False)],
            [make_hat('β', '1'), make_mr(' − ', italic=False), make_hat('β', '2')]
        )
    ]
    
    omath.append(make_frac(num, den))
    return omath

def make_sup_on_element(base_element, sup_text):
    """Create superscript on an existing element"""
    sSup = etree.Element(m('sSup'))
    e = etree.SubElement(sSup, m('e'))
    e.append(base_element)
    sup = etree.SubElement(sSup, m('sup'))
    sup.append(make_mr(sup_text, italic=False))
    return sSup

def build_formula_7():
    """R_t = α + β₁ · Target_t + β₂ · Path_t + β₃ · S_t + β₄ · (S_t × FG_t) + ε_t  (H3)"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('R', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '3'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '4'))
    omath.append(make_mr(' · ', italic=False))
    # (S_t × FG_t)
    omath.append(make_d([
        make_sub('S', 't'),
        make_mr(' × ', italic=False),
        make_sub('FG', 't'),
    ]))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_8():
    """S_t = α + ρS_{t-1} + β₁ · Target_t + β₂ · Path_t + ε_t  (Dynamic)"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_mr('ρ'))
    omath.append(make_sub('S', 't−1'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

# ─── Create oMathPara wrapper ───

def make_omath_para(omath, eq_number=None):
    """Wrap oMath in oMathPara with center alignment and optional equation number"""
    omathPara = etree.Element(m('oMathPara'))
    omathParaPr = etree.SubElement(omathPara, m('oMathParaPr'))
    jc = etree.SubElement(omathParaPr, m('jc'))
    jc.set(m('val'), 'center')
    omathPara.append(omath)
    return omathPara

# ─── Replace paragraph content ───

def replace_formula_paragraph(para, omath_para, eq_number=None):
    """Replace a paragraph's content with an oMathPara"""
    elem = para._element
    
    # Keep pPr
    pPr = elem.find(w('pPr'))
    
    # Remove all existing children except pPr
    for child in list(elem):
        if child.tag != w('pPr'):
            elem.remove(child)
    
    # Add oMathPara
    elem.append(omath_para)
    
    # Add equation number as a separate run (right-aligned tab)
    if eq_number:
        r = etree.SubElement(elem, w('r'))
        rPr = etree.SubElement(r, w('rPr'))
        rFonts = etree.SubElement(rPr, w('rFonts'))
        rFonts.set(w('ascii'), 'Times New Roman')
        rFonts.set(w('hAnsi'), 'Times New Roman')
        t = etree.SubElement(r, w('t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = f'  ({eq_number})'

# ─── Fix inline text ───

def fix_inline_text(doc):
    """Fix LaTeX commands in inline text paragraphs"""
    replacements = [
        # LaTeX commands that should be Greek letters
        (r'\beta_1', 'β₁'),
        (r'\beta_2', 'β₂'),
        (r'\beta_3', 'β₃'),
        (r'\beta_4', 'β₄'),
        (r'\hat\beta_1', 'β̂₁'),
        (r'\hat\beta_2', 'β̂₂'),
        (r'\hat{\beta_1', 'β̂₁'),
        (r'\hat{\beta_2', 'β̂₂'),
        (r'\alpha', 'α'),
        (r'\varepsilon', 'ε'),
        (r'\rho', 'ρ'),
        (r'\cdot', '·'),
        (r'\times', '×'),
        (r'\geq', '≥'),
        (r'\leq', '≤'),
        (r'\textVar', 'Var'),
        (r'\textCov', 'Cov'),
        (r'\textTarget', 'Target'),
        (r'\textPath', 'Path'),
        (r'\frac', ''),
    ]
    
    fixed_count = 0
    for para in doc.paragraphs:
        if para.alignment == 1:  # Skip centered (formula) paragraphs
            continue
        for run in para.runs:
            original = run.text
            if not original:
                continue
            modified = original
            for old, new in replacements:
                modified = modified.replace(old, new)
            if modified != original:
                run.text = modified
                fixed_count += 1
    
    return fixed_count

# ─── Main ───

def main():
    input_path = '/home/z/my-project/monetary-policy-lab/docs/Beyond_the_Rate_JMP_v9.docx'
    output_path = '/home/z/my-project/monetary-policy-lab/docs/Beyond_the_Rate_JMP_v9.1.docx'
    
    doc = Document(input_path)
    
    # Define formula paragraph indices and their builders
    formulas = {
        70: (build_formula_1, '1'),   # LM score
        72: (build_formula_2, '2'),   # CB score
        75: (build_formula_3, '3'),   # Combined sentiment
        97: (build_formula_4, '4'),   # H1 regression
        102: (build_formula_5, '5'),  # H2 regression
        108: (build_formula_6, '6'),  # Wald test
        112: (build_formula_7, '7'),  # H3 regression
        244: (build_formula_8, '8'),  # Dynamic regression
    }
    
    # Replace formula paragraphs
    for idx, (builder, eq_num) in formulas.items():
        para = doc.paragraphs[idx]
        omath = builder()
        omath_para = make_omath_para(omath, eq_num)
        replace_formula_paragraph(para, omath_para, eq_num)
        print(f'  Fixed formula at para {idx} → equation ({eq_num})')
    
    # Fix inline text
    fixed = fix_inline_text(doc)
    print(f'  Fixed {fixed} inline text runs')
    
    # Save
    doc.save(output_path)
    print(f'\nSaved to: {output_path}')

if __name__ == '__main__':
    main()
