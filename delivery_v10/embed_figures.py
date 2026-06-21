"""
Embed Figure 2-5 PNGs into the v10.1 docx, replacing [Figure X about here] placeholders.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCX_IN = '/home/z/my-project/monetary-policy-lab/delivery_v10/Words_Beyond_the_Rate_v10.docx'
DOCX_OUT = '/home/z/my-project/monetary-policy-lab/delivery_v10/Words_Beyond_the_Rate_v10.1.docx'
FIG_DIR = '/home/z/my-project/monetary-policy-lab/delivery_v10/figures'

# Figure mapping: placeholder text → (image file, width in inches)
figure_map = {
    '[Figure 2 about here]': ('figure2_sentiment_vs_shocks.png', 6.5),
    '[Figure 3 about here]': ('figure3_asset_returns.png', 6.0),
    '[Figure 4 about here]': ('figure4_sentiment_by_regime.png', 5.5),
    '[Figure 5 about here]': ('figure5_correlation_heatmap.png', 5.5),
}

doc = Document(DOCX_IN)

replaced = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text in figure_map:
        fig_file, width = figure_map[text]
        fig_path = os.path.join(FIG_DIR, fig_file)
        
        if not os.path.exists(fig_path):
            print(f"  ⚠️ Missing: {fig_path}")
            continue
        
        # Clear the placeholder text
        para.clear()
        
        # Add image to the paragraph
        run = para.add_run()
        run.add_picture(fig_path, width=Inches(width))
        
        # Center the image
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        replaced += 1
        print(f"  ✓ Embedded {fig_file} at paragraph {i} (width={width}\")")

doc.save(DOCX_OUT)
print(f"\n✅ Saved: {DOCX_OUT}")
print(f"  Replaced {replaced}/{len(figure_map)} figure placeholders")
