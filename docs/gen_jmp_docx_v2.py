#!/usr/bin/env python3
"""
Generate JMP paper as docx from paper_v9_35page_fixed.md
Target: 35 pages body text, Times New Roman, academic formatting
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from omath_builder import match_formula, make_omath_para, add_eq_number

MD_PATH = "/home/z/my-project/monetary-policy-lab/docs/paper_v9_35page_fixed.md"
OUT_PATH = "/home/z/my-project/monetary-policy-lab/docs/Beyond_the_Rate_JMP_v9.docx"

# ── Read markdown ──
with open(MD_PATH, "r", encoding="utf-8") as f:
    md_text = f.read()

# ── Parse markdown into structured blocks ──
def parse_md(text):
    blocks = []
    lines = text.split("\n")
    i = 0
    in_table = False
    table_headers = None
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if line.strip() == "":
            if in_table and table_headers:
                blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
                table_headers = None
                table_rows = []
                in_table = False
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == "---":
            if in_table and table_headers:
                blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
                table_headers = None
                table_rows = []
                in_table = False
            i += 1
            continue
        
        # Headings
        if line.startswith("### "):
            if in_table and table_headers:
                blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
                table_headers = None
                table_rows = []
                in_table = False
            blocks.append({"type": "h3", "text": line[4:].strip()})
            i += 1
            continue
        
        if line.startswith("## "):
            if in_table and table_headers:
                blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
                table_headers = None
                table_rows = []
                in_table = False
            blocks.append({"type": "h2", "text": line[3:].strip()})
            i += 1
            continue
        
        if line.startswith("# "):
            if in_table and table_headers:
                blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
                table_headers = None
                table_rows = []
                in_table = False
            blocks.append({"type": "h1", "text": line[2:].strip()})
            i += 1
            continue
        
        # Table rows
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            # Check if separator row
            if all(re.match(r'^[-:]+$', c) for c in cells):
                in_table = True
                i += 1
                continue
            if not in_table:
                table_headers = cells
                in_table = True
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_headers:
                blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
                table_headers = None
                table_rows = []
                in_table = False
        
        # Bullet points
        if line.strip().startswith("- "):
            blocks.append({"type": "bullet", "text": line.strip()[2:]})
            i += 1
            continue
        
        # Table/figure placeholders — SKIP for journal submission
        if line.strip().startswith("[Table") or line.strip().startswith("[Figure"):
            # Don't include "[Table N about here]" placeholders in journal submission
            i += 1
            continue
        
        # Note lines (italic)
        if line.strip().startswith("*Note:") or line.strip().startswith("*Keywords:") or line.strip().startswith("*JEL"):
            blocks.append({"type": "note", "text": line.strip().strip("*")})
            i += 1
            continue
        
        # Regular paragraph
        blocks.append({"type": "para", "text": line.strip()})
        i += 1
    
    # Flush any remaining table
    if in_table and table_headers:
        blocks.append({"type": "table", "headers": table_headers, "rows": table_rows})
    
    return blocks


def clean_latex(text):
    """Convert LaTeX markup to plain text suitable for docx"""
    # Phase 1: Greek letters BEFORE stripping $ delimiters
    # File has \alpha (single backslash in file = r'\alpha' in regex)
    # Must handle \beta_1 before \beta (specific before general)
    text = re.sub(r'\\alpha(?![a-zA-Z])', 'α', text)
    text = re.sub(r'\\beta_1(?![0-9])', 'β₁', text)
    text = re.sub(r'\\beta_2(?![0-9])', 'β₂', text)
    text = re.sub(r'\\beta_3(?![0-9])', 'β₃', text)
    text = re.sub(r'\\beta_4(?![0-9])', 'β₄', text)
    text = re.sub(r'\\beta(?![a-zA-Z_])', 'β', text)
    text = re.sub(r'\\varepsilon(?![a-zA-Z])', 'ε', text)
    text = re.sub(r'\\epsilon(?![a-zA-Z])', 'ε', text)
    text = re.sub(r'\\gamma(?![a-zA-Z])', 'γ', text)
    text = re.sub(r'\\delta(?![a-zA-Z])', 'δ', text)
    text = re.sub(r'\\rho(?![a-zA-Z])', 'ρ', text)
    text = re.sub(r'\\sigma(?![a-zA-Z])', 'σ', text)
    text = re.sub(r'\\Sigma(?![a-zA-Z])', 'Σ', text)
    text = re.sub(r'\\chi(?![a-zA-Z])', 'χ', text)
    text = re.sub(r'\\Delta(?![a-zA-Z])', 'Δ', text)
    text = re.sub(r'\\mu(?![a-zA-Z])', 'μ', text)
    text = re.sub(r'\\lambda(?![a-zA-Z])', 'λ', text)
    text = re.sub(r'\\pi(?![a-zA-Z])', 'π', text)
    text = re.sub(r'\\phi(?![a-zA-Z])', 'φ', text)
    text = re.sub(r'\\omega(?![a-zA-Z])', 'ω', text)

    # Phase 2: Replace $...$ and $$...$$ delimiters
    text = re.sub(r'\$\$([^$]+)\$\$', r'\1', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)

    # Phase 3: LaTeX commands
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
    text = re.sub(r'\\hat\{([^}]+)\}', r'\1̂', text)
    text = re.sub(r'\\bar\{([^}]+)\}', r'\1̄', text)
    text = re.sub(r'\\tilde\{([^}]+)\}', r'\1̃', text)
    text = re.sub(r'\\cdot', '·', text)
    text = re.sub(r'\\times', '×', text)
    text = re.sub(r'\\geq', '≥', text)
    text = re.sub(r'\\leq', '≤', text)
    text = re.sub(r'\\neq', '≠', text)
    text = re.sub(r'\\approx', '≈', text)
    text = re.sub(r'\\Rightarrow', '⇒', text)
    text = re.sub(r'\\Leftarrow', '⇐', text)
    text = re.sub(r'\\qquad', '  ', text)
    text = re.sub(r'\\quad', ' ', text)
    text = re.sub(r'\\,', ' ', text)
    text = re.sub(r'\\left(?![a-zA-Z])', '', text)
    text = re.sub(r'\\right(?![a-zA-Z])', '', text)
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)

    # Phase 4: Subscripts and superscripts
    text = re.sub(r'_t(?![a-zA-Z])', 'ₜ', text)
    text = re.sub(r'_{t-1}', 'ₜ₋₁', text)
    text = re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    text = re.sub(r'\^2(?![0-9])', '²', text)

    # Phase 5: Clean up remaining braces
    text = re.sub(r'\{([^}]+)\}', r'\1', text)

    # Phase 6: Markdown formatting
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Phase 7: Special characters
    text = text.replace('&amp;', '&')
    text = text.replace('S\\&P', 'S&P')
    text = text.replace('S\&P', 'S&P')
    text = text.replace('G\\u00fcrkaynak', 'Gürkaynak')
    for code, char in [('u00d7','×'),('u0394','Δ'),('u03b2','β'),('u03b1','α'),
                       ('u03b5','ε'),('u03c1','ρ'),('u03c7','χ'),('u2248','≈'),
                       ('u2192','→'),('u2013','–'),('u2014','—'),('u201c','"'),
                       ('u201d','"')]:
        text = text.replace('\\' + code, char)

    # Phase 8: Remove remaining LaTeX commands (catch-all, LAST)
    text = re.sub(r'\\[a-zA-Z]+', '', text)

    return text.strip()

def add_formatted_runs(paragraph, text, base_size=12, bold=False, italic=False, color=None):
    """Add text with inline formatting (bold/italic) to a paragraph"""
    # Split on bold/italic markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(clean_latex(part[2:-2]))
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.name = 'Times New Roman'
            if color:
                run.font.color.rgb = color
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(clean_latex(part[1:-1]))
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.name = 'Times New Roman'
            if color:
                run.font.color.rgb = color
        else:
            cleaned = clean_latex(part)
            if cleaned:
                run = paragraph.add_run(cleaned)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(base_size)
                run.font.name = 'Times New Roman'
                if color:
                    run.font.color.rgb = color


def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


# ── Create Document ──
doc = Document()

# ── Set default style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15

# ── Configure heading styles ──
for level, (size, before, after, color) in {
    1: (16, 24, 12, RGBColor(0x0F, 0x20, 0x27)),
    2: (14, 18, 8, RGBColor(0x0F, 0x20, 0x27)),
    3: (12, 12, 6, RGBColor(0x0F, 0x20, 0x27)),
}.items():
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = color
    hs.paragraph_format.space_before = Pt(before)
    hs.paragraph_format.space_after = Pt(after)
    hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    hs.paragraph_format.line_spacing = 1.15
    hs.paragraph_format.keep_with_next = True

# ── Page setup ──
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# ── Add page numbers ──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fldChar2)
for r in [run, run2, run3]:
    r.font.size = Pt(10)
    r.font.name = 'Times New Roman'

# ── Parse and generate ──
blocks = parse_md(md_text)

# Track state
is_title = True
is_abstract = False
after_abstract = False
in_references = False
in_appendix = False

for idx, block in enumerate(blocks):
    btype = block["type"]
    
    # ── Title ──
    if btype == "h1" and is_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(clean_latex(block["text"]))
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x0F, 0x20, 0x27)
        is_title = False
        continue
    
    # ── Author ──
    if btype == "para" and block["text"].startswith("**Eileen Zhang**"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("Eileen Zhang")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        continue
    
    # ── Affiliation ──
    if btype == "para" and "Xi'an Jiaotong-Liverpool" in block["text"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(clean_latex(block["text"]))
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        continue
    
    # ── Abstract heading ──
    if btype == "h2" and block["text"] == "Abstract":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("Abstract")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        is_abstract = True
        continue
    
    # ── Abstract text ──
    if is_abstract and btype == "para":
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        add_formatted_runs(p, block["text"], base_size=11)
        is_abstract = False
        after_abstract = True
        continue
    
    # ── Keywords / JEL ──
    if btype == "note" and ("Keywords" in block["text"] or "JEL" in block["text"]):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(clean_latex(block["text"]))
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        continue
    
    # ── Section headings (h2 = Section level) ──
    if btype == "h2":
        text = block["text"]
        if text == "References":
            in_references = True
            in_appendix = False
        elif text.startswith("Appendix"):
            in_references = False
            in_appendix = True
        else:
            in_references = False
            in_appendix = False
        
        p = doc.add_heading(clean_latex(text), level=1)
        continue
    
    # ── Subsection headings (h3) ──
    if btype == "h3":
        p = doc.add_heading(clean_latex(block["text"]), level=2)
        continue
    
    # ── References ──
    if in_references and btype == "para":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.0
        add_formatted_runs(p, block["text"], base_size=11)
        continue
    
    # ── Tables ──
    if btype == "table":
        headers = [clean_latex(h) for h in block["headers"]]
        rows = block["rows"]
        ncols = len(headers)
        
        # Table caption (look back for **Table N:** pattern)
        table_obj = doc.add_table(rows=1 + len(rows), cols=ncols)
        table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_obj.style = 'Table Grid'
        
        # Header row
        for j, h in enumerate(headers):
            cell = table_obj.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "0F2027")
        
        # Data rows
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci >= ncols:
                    break
                cell = table_obj.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cleaned = clean_latex(val)
                run = p.add_run(cleaned)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                if ri % 2 == 0:
                    set_cell_shading(cell, "F0F4F8")
        
        # Set column widths equally
        for row in table_obj.rows:
            for cell in row.cells:
                cell.width = Inches(6.0 / ncols)
        
        # Add spacing after table
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        continue
    
    # ── Table/Figure placeholders ──
    if btype == "placeholder":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(clean_latex(block["text"]))
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x4A, 0x65, 0x75)
        continue
    
    # ── Note lines ──
    if btype == "note":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_formatted_runs(p, block["text"], base_size=10, italic=True)
        continue
    
    # ── Bullet points ──
    if btype == "bullet":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run("• ")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        add_formatted_runs(p, block["text"], base_size=12)
        continue
    
    # ── Formula paragraphs (lines starting with $$) ──
    if btype == "para" and block["text"].startswith("$$"):
        formula = block["text"].strip("$").strip()
        matched = match_formula(formula)
        if matched:
            builder, eq_num = matched
            omath = builder()
            omath_para = make_omath_para(omath)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p._element.append(omath_para)
            add_eq_number(p._element, eq_num)
        else:
            # Fallback: plain text for unrecognized formulas
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_latex(formula))
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        continue
    
    # ── Regular paragraphs ──
    if btype == "para":
        text = block["text"]
        
        # Table caption lines (e.g., **Table 1: ...)
        if text.startswith("**Table") and ":**" in text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            # Extract table number and title
            caption = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            run = p.add_run(clean_latex(caption))
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            continue
        
        # Formula-like lines
        if text.startswith("$$") and text.endswith("$$"):
            formula = text.strip("$").strip()
            matched = match_formula(formula)
            if matched:
                builder, eq_num = matched
                omath = builder()
                omath_para = make_omath_para(omath)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p._element.append(omath_para)
                add_eq_number(p._element, eq_num)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(clean_latex(formula))
                run.italic = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
            continue
        
        # Data Availability
        if text.startswith("**Data Availability"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            cleaned = clean_latex(text.replace("**Data Availability.**", "Data Availability. "))
            run = p.add_run("Data Availability. ")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            rest = text.replace("**Data Availability.**", "").strip()
            add_formatted_runs(p, rest, base_size=12)
            continue
        
        # Bold-start paragraphs (like **Explanation 1:...)
        if text.startswith("**") and "**" in text[2:]:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.15
            add_formatted_runs(p, text, base_size=12)
            continue
        
        # Regular body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        add_formatted_runs(p, text, base_size=12)
        continue

# ── Save ──
doc.save(OUT_PATH)
print(f"Document saved to {OUT_PATH}")

# ── Verify page count ──
import subprocess
result = subprocess.run(
    ["libreoffice", "--headless", "--convert-to", "pdf", OUT_PATH, "--outdir", "/tmp/"],
    capture_output=True, text=True, timeout=60
)
pdf_path = OUT_PATH.replace(".docx", ".pdf")
tmp_pdf = "/tmp/" + os.path.basename(pdf_path)
if os.path.exists(tmp_pdf):
    import fitz
    pdf_doc = fitz.open(tmp_pdf)
    print(f"Page count: {pdf_doc.page_count}")
    pdf_doc.close()
    os.remove(tmp_pdf)
else:
    print("Could not convert to PDF for page count verification")
