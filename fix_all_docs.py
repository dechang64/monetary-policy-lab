#!/usr/bin/env python3
"""
Fix all JMP docx files:
1. Convert broken formula text to proper Word oMath objects
2. Remove [Table/Figure about here] placeholders
3. Fix inline text with missing β symbols
"""

from docx import Document
from docx.shared import Pt
from lxml import etree
import re
import sys
import os

# Namespaces
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def m(tag): return f'{{{M_NS}}}{tag}'
def w(tag): return f'{{{W_NS}}}{tag}'

# ─── oMath Builder Helpers ───

def make_mr(text, italic=True):
    r = etree.Element(m('r'))
    rPr = etree.SubElement(r, m('rPr'))
    if italic:
        sty = etree.SubElement(rPr, m('sty'))
        sty.set(m('val'), 'i')
    t = etree.SubElement(r, m('t'))
    t.text = text
    return r

def make_sub(base_text, sub_text, base_italic=True):
    sSub = etree.Element(m('sSub'))
    e = etree.SubElement(sSub, m('e'))
    e.append(make_mr(base_text, italic=base_italic))
    sub = etree.SubElement(sSub, m('sub'))
    sub.append(make_mr(sub_text, italic=False))
    return sSub

def make_sup_on_element(base_element, sup_text):
    sSup = etree.Element(m('sSup'))
    e = etree.SubElement(sSup, m('e'))
    e.append(base_element)
    sup = etree.SubElement(sSup, m('sup'))
    sup.append(make_mr(sup_text, italic=False))
    return sSup

def make_frac(num_children, den_children):
    f = etree.Element(m('f'))
    fPr = etree.SubElement(f, m('fPr'))
    ctrlPr = etree.SubElement(fPr, m('ctrlPr'))
    num = etree.SubElement(f, m('num'))
    for child in num_children:
        num.append(child)
    den = etree.SubElement(f, m('den'))
    for child in den_children:
        den.append(child)
    return f

def make_hat(base_text, sub_text=None):
    acc = etree.Element(m('acc'))
    accPr = etree.SubElement(acc, m('accPr'))
    chr_elem = etree.SubElement(accPr, m('chr'))
    chr_elem.set(m('val'), '̂')
    ctrlPr = etree.SubElement(accPr, m('ctrlPr'))
    e = etree.SubElement(acc, m('e'))
    if sub_text:
        e.append(make_sub(base_text, sub_text))
    else:
        e.append(make_mr(base_text))
    return acc

def make_d(children):
    d = etree.Element(m('d'))
    dPr = etree.SubElement(d, m('dPr'))
    ctrlPr = etree.SubElement(dPr, m('ctrlPr'))
    e = etree.SubElement(d, m('e'))
    for child in children:
        e.append(child)
    return d

def make_func(name_children, arg_children):
    func = etree.Element(m('func'))
    funcPr = etree.SubElement(func, m('funcPr'))
    ctrlPr = etree.SubElement(funcPr, m('ctrlPr'))
    fName = etree.SubElement(func, m('fName'))
    for child in name_children:
        fName.append(child)
    e = etree.SubElement(func, m('e'))
    for child in arg_children:
        e.append(child)
    return func

# ─── Formula Builders ───

def build_formula_1():
    """LM_t = (Positive words_t - Negative words_t) / Total words_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('LM', 't'))
    omath.append(make_mr(' = ', italic=False))
    num = [make_mr('Positive wordsₜ − Negative wordsₜ', italic=False)]
    den = [make_mr('Total wordsₜ', italic=False)]
    omath.append(make_frac(num, den))
    return omath

def build_formula_2():
    """CB_t = (Hawkish words_t - Dovish words_t) / Total words_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('CB', 't'))
    omath.append(make_mr(' = ', italic=False))
    num = [make_mr('Hawkish wordsₜ − Dovish wordsₜ', italic=False)]
    den = [make_mr('Total wordsₜ', italic=False)]
    omath.append(make_frac(num, den))
    return omath

def build_formula_3():
    """S_t = 0.5 × LM_t + 0.5 × CB_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = 0.5 × ', italic=False))
    omath.append(make_sub('LM', 't'))
    omath.append(make_mr(' + 0.5 × ', italic=False))
    omath.append(make_sub('CB', 't'))
    return omath

def build_formula_4():
    """S_t = α + β₁ · Target_t + β₂ · Path_t + ε_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_5():
    """R_t = α + β₁ · Target_t + β₂ · Path_t + ε_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('R', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_6():
    """W = (β̂₁ − β̂₂)² / Var(β̂₁ − β̂₂)"""
    omath = etree.Element(m('oMath'))
    omath.append(make_mr('W'))
    omath.append(make_mr(' = ', italic=False))
    paren_content = make_d([
        make_hat('β', '1'),
        make_mr(' − ', italic=False),
        make_hat('β', '2'),
    ])
    num = [make_sup_on_element(paren_content, '2')]
    den = [
        make_func(
            [make_mr('Var', italic=False)],
            [make_hat('β', '1'), make_mr(' − ', italic=False), make_hat('β', '2')]
        )
    ]
    omath.append(make_frac(num, den))
    return omath

def build_formula_7():
    """R_t = α + β₁ · Target_t + β₂ · Path_t + β₃ · S_t + β₄ · (S_t × FG_t) + ε_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('R', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '3'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '4'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_d([
        make_sub('S', 't'),
        make_mr(' × ', italic=False),
        make_sub('FG', 't'),
    ]))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

def build_formula_8():
    """S_t = α + ρS_{t-1} + β₁ · Target_t + β₂ · Path_t + ε_t"""
    omath = etree.Element(m('oMath'))
    omath.append(make_sub('S', 't'))
    omath.append(make_mr(' = ', italic=False))
    omath.append(make_mr('α'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_mr('ρ'))
    omath.append(make_sub('S', 't−1'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '1'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Target', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('β', '2'))
    omath.append(make_mr(' · ', italic=False))
    omath.append(make_sub('Path', 't'))
    omath.append(make_mr(' + ', italic=False))
    omath.append(make_sub('ε', 't'))
    return omath

FORMULA_BUILDERS = [
    (build_formula_1, '1'),
    (build_formula_2, '2'),
    (build_formula_3, '3'),
    (build_formula_4, '4'),
    (build_formula_5, '5'),
    (build_formula_6, '6'),
    (build_formula_7, '7'),
    (build_formula_8, '8'),
]

# ─── Formula detection patterns ───
# Each pattern matches a specific broken formula text

FORMULA_PATTERNS = [
    # (pattern, builder_index)
    (r'LMₜ\s*=\s*Positive', 0),           # Formula 1: LM
    (r'CBₜ\s*=\s*Hawkish', 1),             # Formula 2: CB
    (r'Sₜ\s*=\s*0\.5\s*×\s*LM', 2),       # Formula 3: Combined
    (r'Sₜ\s*=\s*.*β₁.*Target.*β₂.*Path.*ₜ', 3),  # Formula 4: H1 (S_t = ...)
    (r'Rₜ\s*=\s*.*β₁.*Target.*β₂.*Path.*ₜ', 4),  # Formula 5: H2 (R_t = ...)
    (r'W\s*=\s*.*̂.*_1.*̂.*_2', 5),        # Formula 6: Wald
    (r'Rₜ\s*=\s*.*β₁.*β₂.*β₃.*β₄', 6),   # Formula 7: H3
    (r'Sₜ\s*=\s*.*ρ.*Sₜ', 7),             # Formula 8: Dynamic
]

# Also match LaTeX versions
FORMULA_PATTERNS_LATEX = [
    (r'\\textLMₜ\s*=\s*\\frac', 0),
    (r'\\textCBₜ\s*=\s*\\frac', 1),
    (r'Sₜ\s*=\s*0\.5\s*\\times', 2),
    (r'Sₜ\s*=\s*\\alpha\s*\+\\s*\\beta_1', 3),
    (r'Rₜ\s*=\s*\\alpha\s*\+\\s*\\beta_1', 4),
    (r'W\s*=\s*\\frac', 5),
    (r'Rₜ\s*=\s*\\alpha.*\\beta_3', 6),
    (r'Sₜ\s*=\s*\\alpha\s*\+\\s*\\rho', 7),
]


def make_omath_para(omath):
    omathPara = etree.Element(m('oMathPara'))
    omathParaPr = etree.SubElement(omathPara, m('oMathParaPr'))
    jc = etree.SubElement(omathParaPr, m('jc'))
    jc.set(m('val'), 'center')
    omathPara.append(omath)
    return omathPara


def replace_formula_paragraph(para, omath_para, eq_number):
    elem = para._element
    pPr = elem.find(w('pPr'))
    for child in list(elem):
        if child.tag != w('pPr'):
            elem.remove(child)
    elem.append(omath_para)
    # Add equation number
    r = etree.SubElement(elem, w('r'))
    rPr = etree.SubElement(r, w('rPr'))
    rFonts = etree.SubElement(rPr, w('rFonts'))
    rFonts.set(w('ascii'), 'Times New Roman')
    rFonts.set(w('hAnsi'), 'Times New Roman')
    t = etree.SubElement(r, w('t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = f'  ({eq_number})'


def fix_file(input_path, output_path=None):
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f'{base}_fixed{ext}'
    
    doc = Document(input_path)
    name = os.path.basename(input_path)
    
    # ── Step 1: Find and fix formula paragraphs ──
    formula_fixes = 0
    used_builders = set()
    
    for i, para in enumerate(doc.paragraphs):
        if para.alignment != 1:  # Only centered paragraphs
            continue
        text = para.text.strip()
        if not text:
            continue
        
        # Try each pattern
        for pattern, builder_idx in FORMULA_PATTERNS + FORMULA_PATTERNS_LATEX:
            if re.search(pattern, text):
                builder, eq_num = FORMULA_BUILDERS[builder_idx]
                if builder_idx not in used_builders:
                    omath = builder()
                    omath_para = make_omath_para(omath)
                    replace_formula_paragraph(para, omath_para, eq_num)
                    used_builders.add(builder_idx)
                    formula_fixes += 1
                    print(f'  ✓ Formula ({eq_num}) at para {i}')
                break
    
    # ── Step 2: Remove [Table/Figure about here] placeholders ──
    placeholder_count = 0
    paras_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('[Table') or text.startswith('[Figure'):
            if 'about here' in text:
                paras_to_remove.append(para)
                placeholder_count += 1
    
    for para in paras_to_remove:
        parent = para._element.getparent()
        if parent is not None:
            parent.remove(para._element)
    
    if placeholder_count:
        print(f'  ✓ Removed {placeholder_count} placeholders')
    
    # ── Step 3: Fix inline text ──
    inline_fixes = 0
    replacements = [
        # LaTeX → Unicode
        (r'\beta_1', 'β₁'),
        (r'\beta_2', 'β₂'),
        (r'\beta_3', 'β₃'),
        (r'\beta_4', 'β₄'),
        (r'\hatβ₁', 'β̂₁'),
        (r'\hatβ₂', 'β̂₂'),
        (r'\hat{\beta_1', 'β̂₁'),
        (r'\hat{\beta_2', 'β̂₂'),
        (r'\hat\beta_1', 'β̂₁'),
        (r'\hat\beta_2', 'β̂₂'),
        (r'\alpha', 'α'),
        (r'\varepsilon', 'ε'),
        (r'\rho', 'ρ'),
        (r'\cdot', '·'),
        (r'\times', '×'),
        (r'\geq', '≥'),
        (r'\leq', '≤'),
        (r'\textVar', 'Var'),
        (r'\textCov', 'Cov'),
        (r'\textTarget', 'Target'),
        (r'\textPath', 'Path'),
        (r'\frac', ''),
        # Also fix broken Unicode patterns
        ('̂_1', 'β̂₁'),  # hat + _1 → β̂₁
        ('̂_2', 'β̂₂'),  # hat + _2 → β̂₂
    ]
    
    for para in doc.paragraphs:
        if para.alignment == 1:  # Skip centered (formula) paragraphs
            continue
        for run in para.runs:
            if not run.text:
                continue
            original = run.text
            modified = original
            for old, new in replacements:
                modified = modified.replace(old, new)
            if modified != original:
                run.text = modified
                inline_fixes += 1
    
    if inline_fixes:
        print(f'  ✓ Fixed {inline_fixes} inline text runs')
    
    # ── Save ──
    doc.save(output_path)
    print(f'  → Saved: {output_path}')
    print(f'  Total: {formula_fixes} formulas, {placeholder_count} placeholders, {inline_fixes} inline fixes')
    return formula_fixes + placeholder_count + inline_fixes


def main():
    files_to_fix = [
        '/home/z/my-project/monetary-policy-lab/delivery_v9.9/Beyond_the_Rate_JMP_v9.docx',
        '/home/z/my-project/monetary-policy-lab/docs/Beyond_the_Rate_JMP_v9.docx',
        '/home/z/my-project/monetary-policy-lab/docs/paper_v9_35page.docx',
    ]
    
    total = 0
    for fpath in files_to_fix:
        name = os.path.basename(fpath)
        print(f'\n{"="*60}')
        print(f'Fixing: {name}')
        print(f'{"="*60}')
        try:
            total += fix_file(fpath)
        except Exception as e:
            print(f'  ERROR: {e}')
    
    print(f'\n{"="*60}')
    print(f'All done. Total fixes: {total}')

if __name__ == '__main__':
    main()
