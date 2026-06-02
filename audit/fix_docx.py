"""
Comprehensive fix script for Beyond the Rate v9.2 docx
Fixes: figure order, wrong numbers, missing table content
"""
from docx import Document
from docx.shared import Inches, Pt
from copy import deepcopy
import re

INPUT = '/home/z/my-project/upload/6a1b847f73da0208f1c22948_beyond_the_rate_jmp_v9-2_论文.docx'
OUTPUT = '/home/z/my-project/monetary-policy-lab/delivery_v9.2/Beyond_the_Rate_JMP_v9.2_论文_审计修正版.docx'

doc = Document(INPUT)

# ===== Fix 1: p = 0.726 → 0.526 (6 occurrences) =====
count_726 = 0
for p in doc.paragraphs:
    if '0.726' in p.text:
        for run in p.runs:
            if '0.726' in run.text:
                run.text = run.text.replace('0.726', '0.526')
                count_726 += 1
print(f"✓ Fixed p=0.726 → 0.526: {count_726} runs")

# Also fix in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if '0.726' in cell.text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '0.726' in run.text:
                            run.text = run.text.replace('0.726', '0.526')
                            count_726 += 1
print(f"✓ Total p=0.726 fixes (including tables): {count_726}")

# ===== Fix 2: Figure numbering =====
# Current order: Fig2(133), Fig1(147), Fig3(168), Fig5(183), Fig4(236), Fig6(244), Fig7(250)
# Correct order: Fig1, Fig2, Fig3, Fig4, Fig5, Fig6, Fig7
# 
# Strategy: Rename figures to match their appearance order
# Para 133: "Figure 2" → "Figure 2" (keep, will be renumbered)
# Para 147: "Figure 1" → needs to become Figure 1
# Para 183: "Figure 5" → needs to become Figure 4
# Para 236: "Figure 4" → needs to become Figure 5

# The correct numbering based on appearance order:
# Para 133: Figure 2 → Figure 2 (2nd figure, but appears first → should be Figure 1)
# Wait, let me think about this differently.
# 
# The figures should be numbered by their appearance order in the text:
# 1st appearance: Para 133 (currently "Figure 2") → should be Figure 1
# 2nd appearance: Para 147 (currently "Figure 1") → should be Figure 2
# 3rd appearance: Para 168 (currently "Figure 3") → should be Figure 3
# 4th appearance: Para 183 (currently "Figure 5") → should be Figure 4
# 5th appearance: Para 236 (currently "Figure 4") → should be Figure 5
# 6th appearance: Para 244 (currently "Figure 6") → should be Figure 6
# 7th appearance: Para 250 (currently "Figure 7") → should be Figure 7

figure_renames = {
    'Figure 2: Sentiment vs. Target and Path Shocks': 'Figure 1: Sentiment vs. Target and Path Shocks',
    'Figure 1: FOMC Statement Sentiment and Monetary Policy Shocks': 'Figure 2: FOMC Statement Sentiment and Monetary Policy Shocks',
    'Figure 5: Sentiment by Decision Type': 'Figure 4: Sentiment by Decision Type',
    'Figure 4: Sentiment Measure Comparison': 'Figure 5: Sentiment Measure Comparison',
}

count_fig = 0
for p in doc.paragraphs:
    for old, new in figure_renames.items():
        if old in p.text:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count_fig += 1
                    print(f"  Renamed: {old[:40]}... → {new[:40]}...")

# Also fix in-text references like "Figure 2 shows" → "Figure 1 shows"
# This is trickier - need to update all cross-references
# For now, let me fix the common patterns

# Fix cross-references in text
text_renames = {
    'Figure 2 shows': 'Figure 1 shows',
    'Figure 1 shows': 'Figure 2 shows',  # This will be applied AFTER the first rename
    'Figure 5 shows': 'Figure 4 shows',
    'Figure 4 shows': 'Figure 5 shows',
}

# Actually, this is dangerous because of ordering. Let me use a different approach.
# First pass: rename to temporary names, second pass: rename to final names
temp_renames = {
    'Figure 2': 'FIG_TEMP_A',
    'Figure 1': 'FIG_TEMP_B',
    'Figure 5': 'FIG_TEMP_C',
    'Figure 4': 'FIG_TEMP_D',
}

final_renames = {
    'FIG_TEMP_A': 'Figure 1',
    'FIG_TEMP_B': 'Figure 2',
    'FIG_TEMP_C': 'Figure 4',
    'FIG_TEMP_D': 'Figure 5',
}

# Apply temp renames to all paragraphs
for p in doc.paragraphs:
    for run in p.runs:
        for old, new in temp_renames.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

# Apply final renames
for p in doc.paragraphs:
    for run in p.runs:
        for old, new in final_renames.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

# Also fix in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    for old, new in temp_renames.items():
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                    for old, new in final_renames.items():
                        if old in run.text:
                            run.text = run.text.replace(old, new)

print(f"✓ Fixed figure numbering")

# ===== Fix 3: Table A2 Lag Sensitivity =====
# Find Table A2 and fix the numbers
for i, table in enumerate(doc.tables):
    # Check if this is Table A2
    first_cell = table.rows[0].cells[0].text.strip() if table.rows else ''
    if 'Lag' in first_cell and 't-stat' in table.rows[0].cells[1].text:
        print(f"\nFound Table A2 at table index {i}")
        # Fix the values
        correct_values = {
            '1': {'t': '2.17', 'p': '0.032'},
            '2': {'t': '2.24', 'p': '0.027'},
            '4': {'t': '2.43', 'p': '0.017'},
            '6': {'t': '2.56', 'p': '0.012'},
        }
        for row in table.rows[1:]:  # Skip header
            lag_val = row.cells[0].text.strip()
            if lag_val in correct_values:
                cv = correct_values[lag_val]
                # Fix t-stat
                for p in row.cells[1].paragraphs:
                    for run in p.runs:
                        run.text = cv['t']
                # Fix p-value
                for p in row.cells[2].paragraphs:
                    for run in p.runs:
                        run.text = cv['p']
                print(f"  Fixed lag={lag_val}: t={cv['t']}, p={cv['p']}")

# ===== Fix 4: Table 6 Regime Analysis =====
for i, table in enumerate(doc.tables):
    first_cell = table.rows[0].cells[0].text.strip() if table.rows else ''
    if 'Regime' in first_cell or 'Decision' in first_cell:
        # Check if this has rate_hike, rate_cut rows
        for row in table.rows:
            cell0 = row.cells[0].text.strip()
            if 'Rate hike' in cell0 or 'Hike' in cell0:
                # Fix β_T p: 0.013 → 0.026
                for cell in row.cells:
                    if '0.013' in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace('0.013', '0.026')
                        print(f"  Fixed Rate hike β_T p: 0.013 → 0.026")
            elif 'Rate cut' in cell0 or 'Cut' in cell0:
                # Fix β_P p: <0.001 → 0.006
                for cell in row.cells:
                    if '<0.001' in cell.text or '< 0.001' in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace('<0.001', '0.006').replace('< 0.001', '0.006')
                        print(f"  Fixed Rate cut β_P p: <0.001 → 0.006")
                    elif '0.089' in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace('0.089', '0.128')
                        print(f"  Fixed Rate cut β_T p: 0.089 → 0.128")

# ===== Fix 5: Table 9 Subsample =====
for i, table in enumerate(doc.tables):
    for row in table.rows:
        cell0 = row.cells[0].text.strip()
        if 'Post-crisis' in cell0:
            # Fix R²: 1.6% → 0.6%
            for cell in row.cells:
                if '1.6%' in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace('1.6%', '0.6%')
                    print(f"  Fixed Post-crisis R²: 1.6% → 0.6%")
                if '0.056' in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace('0.056', '0.258')
                    print(f"  Fixed Post-crisis β_P p: 0.056 → 0.258")
                if '0.699' in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace('0.699', '0.633')
                    print(f"  Fixed Post-crisis β_T p: 0.699 → 0.633")

# ===== Fix 6: Table 8 Minutes CB β_P =====
for i, table in enumerate(doc.tables):
    for row in table.rows:
        cell0 = row.cells[0].text.strip()
        if 'Minutes CB' in cell0:
            for cell in row.cells:
                if '0.002876' in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = run.text.replace('0.002876', '0.001423')
                    print(f"  Fixed Minutes CB β_P: 0.002876 → 0.001423")

# ===== Fix 7: "undetectable" → "statistically insignificant" =====
for p in doc.paragraphs:
    if 'undetectable' in p.text:
        for run in p.runs:
            if 'undetectable' in run.text:
                run.text = run.text.replace('undetectable', 'statistically insignificant')
                print(f"  Fixed 'undetectable' → 'statistically insignificant'")

# ===== Fix 8: Table 7 content =====
# Find Table 7 title paragraph and add the missing table content after it
for i, p in enumerate(doc.paragraphs):
    if 'Table 7: Sentiment Dictionary Comparison' in p.text:
        print(f"\nFound Table 7 title at paragraph {i}")
        # Check if the next paragraph has the table content
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.strip()
            if not next_text or next_text == '[]':
                # Need to add table content
                # Insert a new paragraph with the table data
                # python-docx doesn't support inserting paragraphs at specific positions easily
                # Let me modify the next paragraph instead
                doc.paragraphs[i+1].clear()
                run = doc.paragraphs[i+1].add_run(
                    'Combined (LM + CB)\t1.57%\t0.000577 (0.017)\t0.000633 (0.152)\t117\n'
                    'LM only\t0.33%\t0.000288 (0.476)\t0.000465 (0.553)\t117\n'
                    'CB only\t3.90%\t0.000865 (<0.001)\t0.000800 (0.033)\t117'
                )
                run.font.size = Pt(10)
                print(f"  Added Table 7 content")

# Save
doc.save(OUTPUT)
print(f"\n✅ Saved corrected document to: {OUTPUT}")

