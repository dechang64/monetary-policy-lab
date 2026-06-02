#!/usr/bin/env python3
"""
Generate v10.2 docx matching v9.2 formatting quality.
- Times New Roman 12pt
- 1.15x line spacing (276 twips)
- 1.25in left/right margins, 1in top/bottom
- Equations as Unicode text with (1)-(8) numbering
- Tables with proper formatting
- Figures embedded
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Paths
BASE = '/home/z/my-project/monetary-policy-lab'
MD_PATH = os.path.join(BASE, 'delivery_v10', 'Words_Beyond_the_Rate_v10_source.md')
FIG_DIR = os.path.join(BASE, 'delivery_v10', 'figures')
OUT_PATH = os.path.join(BASE, 'delivery_v10', 'Words_Beyond_the_Rate_v10.2.docx')

# Read markdown
with open(MD_PATH, 'r', encoding='utf-8') as f:
    md_text = f.read()

# Equation counter
eq_counter = 0

# Equation mapping: LaTeX -> Unicode text with numbering
EQUATION_MAP = {
    r'$$\text{LM}_t = \frac{\text{Positive words}_t - \text{Negative words}_t}{\text{Total words}_t}$$':
        'LMₜ = (Positive wordsₜ − Negative wordsₜ) / Total wordsₜ    (1)',
    r'$$\text{CB}_t = \frac{\text{Hawkish words}_t - \text{Dovish words}_t}{\text{Total words}_t}$$':
        'CBₜ = (Hawkish wordsₜ − Dovish wordsₜ) / Total wordsₜ    (2)',
    r'$$S_t = 0.5 \times \text{LM}_t + 0.5 \times \text{CB}_t$$':
        'Sₜ = 0.5 × LMₜ + 0.5 × CBₜ    (3)',
    r'$$S_t = \alpha + \beta_1 \cdot \text{Target}_t + \beta_2 \cdot \text{Path}_t + \varepsilon_t$$':
        'Sₜ = α + β₁ · Targetₜ + β₂ · Pathₜ + εₜ    (4)',
    r'$$R_t = \alpha + \beta_1 \cdot \text{Target}_t + \beta_2 \cdot \text{Path}_t + \varepsilon_t$$':
        'Rₜ = α + β₁ · Targetₜ + β₂ · Pathₜ + εₜ    (5)',
    r'$$W = \frac{(\hat{\beta}_1 - \hat{\beta}_2)^2}{\text{Var}(\hat{\beta}_1 - \hat{\beta}_2)}$$':
        'W = (β̂₁ − β̂₂)² / Var(β̂₁ − β̂₂)    (6)',
    r'$$R_t = \alpha + \beta_1 \cdot \text{Target}_t + \beta_2 \cdot \text{Path}_t + \beta_3 \cdot S_t + \beta_4 \cdot (S_t \times FG_t) + \varepsilon_t$$':
        'Rₜ = α + β₁ · Targetₜ + β₂ · Pathₜ + β₃ · Sₜ + β₄ · (Sₜ × FGₜ) + εₜ    (7)',
    r'$$S_t = \alpha + \rho S_{t-1} + \beta_1 \cdot \text{Target}_t + \beta_2 \cdot \text{Path}_t + \varepsilon_t$$':
        'Sₜ = α + ρ · Sₜ₋₁ + β₁ · Targetₜ + β₂ · Pathₜ + εₜ    (8)',
}

# Figure mapping
FIGURE_MAP = {
    'figures/figure1_framework.png': 'figure1_framework.png',
    'figures/figure2_sentiment_vs_shocks.png': 'figure2_sentiment_vs_shocks.png',
    'figures/figure3_asset_returns.png': 'figure3_asset_returns.png',
    'figures/figure4_sentiment_by_regime.png': 'figure4_sentiment_by_regime.png',
    'figures/figure5_correlation_heatmap.png': 'figure5_correlation_heatmap.png',
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_table(doc, header_row, data_rows, note=None):
    """Add a properly formatted table."""
    ncols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for j, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        set_cell_shading(cell, 'E8E8E8')
    
    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    
    # Add note if provided
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(note.strip())
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
    
    return table

def parse_markdown_table(lines):
    """Parse markdown table lines into header and data rows."""
    header = [c.strip() for c in lines[0].split('|')[1:-1]]
    # Skip separator line (line 1)
    data = []
    for line in lines[2:]:
        if line.strip() and not all(c in '|-: ' for c in line.strip()):
            row = [c.strip() for c in line.split('|')[1:-1]]
            data.append(row)
    return header, data

def add_body_paragraph(doc, text, first_para=False):
    """Add a body text paragraph with proper formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    if first_para:
        p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Handle inline formatting
    # Bold text **text**
    # Italic text *text*
    # Inline math $...$ -> Unicode
    
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        elif part.startswith('$') and part.endswith('$'):
            # Inline math - convert to Unicode
            math_text = part[1:-1]
            math_text = math_text.replace('\\beta', 'β').replace('\\alpha', 'α')
            math_text = math_text.replace('\\varepsilon', 'ε').replace('\\rho', 'ρ')
            math_text = math_text.replace('\\hat{\\beta}', 'β̂').replace('\\hat{\\beta}_1', 'β̂₁')
            math_text = math_text.replace('\\text{', '').replace('}', '')
            math_text = math_text.replace('_t', 'ₜ').replace('_1', '₁').replace('_2', '₂')
            math_text = math_text.replace('_3', '₃').replace('_4', '₄')
            math_text = math_text.replace('\\times', '×').replace('\\cdot', '·')
            math_text = math_text.replace('\\chi^2', 'χ²')
            run = p.add_run(math_text)
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        else:
            run = p.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    return p

def add_equation_paragraph(doc, eq_text):
    """Add a centered equation paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(eq_text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def add_heading(doc, text, level):
    """Add a heading with proper formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12) if level > 1 else Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    return p

def add_figure(doc, fig_path, caption, width=Inches(5.5)):
    """Add a figure with caption."""
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run()
        run.add_picture(fig_path, width=width)
    
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(3)
    p_cap.paragraph_format.space_after = Pt(6)
    run = p_cap.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p_cap

def add_blockquote(doc, text):
    """Add a blockquote (for hypotheses, terminology notes)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    # Handle bold within blockquote
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.italic = True
    return p

# Create document
doc = Document()

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.15

# ===== TITLE =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Words Beyond the Rate: High-Frequency Monetary Policy Shocks and FOMC Language')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

# Author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Eileen Zhang')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Affiliation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Academy of AI, Xi\'an Jiaotong-Liverpool University, Suzhou, China')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.italic = True

# Horizontal rule
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(6)

# ===== PARSE MARKDOWN =====
lines = md_text.split('\n')
i = 0
in_table = False
table_lines = []
table_note = None

# Skip the title and author lines (already added)
skip_until_abstract = True

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # Skip title/author section (already added above)
    if skip_until_abstract:
        if stripped == '## Abstract':
            skip_until_abstract = False
            add_heading(doc, 'Abstract', 2)
            i += 1
            continue
        i += 1
        continue
    
    # Skip horizontal rules
    if stripped == '---':
        i += 1
        continue
    
    # Empty line
    if not stripped:
        i += 1
        continue
    
    # Display equation
    if stripped.startswith('$$') and stripped.endswith('$$'):
        eq_latex = stripped
        if eq_latex in EQUATION_MAP:
            add_equation_paragraph(doc, EQUATION_MAP[eq_latex])
        else:
            # Fallback: try to convert
            eq_text = eq_latex[2:-2]
            eq_text = eq_text.replace('\\text{', '').replace('}', '')
            eq_text = eq_text.replace('\\frac{', '').replace('\\alpha', 'α')
            eq_text = eq_text.replace('\\beta_1', 'β₁').replace('\\beta_2', 'β₂')
            eq_text = eq_text.replace('\\beta_3', 'β₃').replace('\\beta_4', 'β₄')
            eq_text = eq_text.replace('\\varepsilon', 'ε').replace('\\rho', 'ρ')
            eq_text = eq_text.replace('\\hat{\\beta}_1', 'β̂₁').replace('\\hat{\\beta}_2', 'β̂₂')
            eq_text = eq_text.replace('\\times', '×').replace('\\cdot', '·')
            eq_text = eq_text.replace('_t', 'ₜ').replace('_{t-1}', 'ₜ₋₁')
            add_equation_paragraph(doc, eq_text)
        i += 1
        continue
    
    # Multi-line display equation (starts with $$ but doesn't end on same line)
    if stripped.startswith('$$') and not stripped.endswith('$$'):
        eq_lines = [stripped]
        i += 1
        while i < len(lines) and not lines[i].strip().endswith('$$'):
            eq_lines.append(lines[i].strip())
            i += 1
        if i < len(lines):
            eq_lines.append(lines[i].strip())
        eq_latex = ' '.join(eq_lines)
        if eq_latex in EQUATION_MAP:
            add_equation_paragraph(doc, EQUATION_MAP[eq_latex])
        i += 1
        continue
    
    # Headings
    if stripped.startswith('### '):
        add_heading(doc, stripped[4:], 3)
        i += 1
        continue
    if stripped.startswith('## '):
        add_heading(doc, stripped[3:], 2)
        i += 1
        continue
    
    # Figure
    fig_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
    if fig_match:
        caption = fig_match.group(1)
        fig_file = fig_match.group(2)
        fig_name = FIGURE_MAP.get(fig_file, os.path.basename(fig_file))
        fig_path = os.path.join(FIG_DIR, fig_name)
        
        # Determine width based on figure
        if 'framework' in fig_name:
            width = Inches(5.5)
        elif 'sentiment_vs_shocks' in fig_name or 'sentiment_shocks' in fig_name:
            width = Inches(6.0)
        elif 'asset_returns' in fig_name:
            width = Inches(5.5)
        elif 'regime' in fig_name:
            width = Inches(5.5)
        elif 'heatmap' in fig_name or 'correlation' in fig_name:
            width = Inches(5.0)
        else:
            width = Inches(5.5)
        
        add_figure(doc, fig_path, caption, width)
        i += 1
        continue
    
    # Blockquote
    if stripped.startswith('> '):
        bq_text = stripped[2:]
        # Collect multi-line blockquote
        while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
            i += 1
            bq_text += ' ' + lines[i].strip()[2:]
        add_blockquote(doc, bq_text)
        i += 1
        continue
    
    # Table
    if '|' in stripped and stripped.startswith('|'):
        # Collect all table lines
        table_lines = [stripped]
        i += 1
        while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        
        # Check for note after table
        if i < len(lines) and lines[i].strip().startswith('*') and 'Note' in lines[i]:
            table_note = lines[i].strip()
            i += 1
        else:
            table_note = None
        
        # Parse and add table
        header, data = parse_markdown_table(table_lines)
        
        # Clean up markdown formatting in cells
        def clean_cell(text):
            text = text.replace('**', '').replace('*', '')
            # Convert inline math
            text = re.sub(r'\$([^$]+)\$', lambda m: m.group(1).replace('\\beta_T', 'β_T').replace('\\beta_P', 'β_P').replace('\\beta', 'β').replace('\\alpha', 'α'), text)
            return text
        
        header = [clean_cell(h) for h in header]
        data = [[clean_cell(c) for c in row] for row in data]
        
        add_formatted_table(doc, header, data, table_note)
        continue
    
    # Table caption (bold line before table)
    if stripped.startswith('**Table'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        # Remove ** markers
        caption_text = stripped.replace('**', '')
        run = p.add_run(caption_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        i += 1
        continue
    
    # Note line (italic)
    if stripped.startswith('*Note:') or stripped.startswith('*Note.'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        note_text = stripped.strip('*')
        run = p.add_run(note_text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        i += 1
        continue
    
    # Regular paragraph
    add_body_paragraph(doc, stripped, first_para=True)
    i += 1

# Save
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")

# Verify page count with LibreOffice
import subprocess
result = subprocess.run(
    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
     os.path.dirname(OUT_PATH), OUT_PATH],
    capture_output=True, text=True, timeout=120
)
print(f"LibreOffice conversion: {result.returncode}")

# Count PDF pages
import fitz
pdf_path = OUT_PATH.replace('.docx', '.pdf')
if os.path.exists(pdf_path):
    pdf_doc = fitz.open(pdf_path)
    print(f"PDF pages: {pdf_doc.page_count}")
    pdf_doc.close()
